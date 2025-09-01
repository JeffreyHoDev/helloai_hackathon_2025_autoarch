# my_adk_mcp_server.py
import asyncio
import json
import os
import sys
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
from dotenv import load_dotenv


import re
import requests
import base64
from io import BytesIO
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors

import markdown
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration

from fastmcp import FastMCP

# --- Load Environment Variables (If ADK tools need them, e.g., API keys) ---
load_dotenv() # Create a .env file in the same directory if needed

# Document processing libraries
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    raise ImportError("Please install python-docx with: uv pip install python-docx")

try:
    import pandas as pd
    import openpyxl
except ImportError:
    raise ImportError("Please install pandas and openpyxl with: uv pip install pandas openpyxl")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    raise ImportError("Please install reportlab with: uv pip install reportlab")

try:
    import docx2pdf
except ImportError:
    raise ImportError("Please install docx2pdf with: uv pip install docx2pdf")

# Set up logging
log_dir = Path(__file__).parent.parent / "logs"
log_dir.mkdir(exist_ok=True)
log_file = log_dir / "document_mcp.log"

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(log_file)
    ]
)
logger = logging.getLogger(__name__)

mcp = FastMCP("file_mcp")

# --- Prepare the ADK Tool ---
# Instantiate the ADK tool you want to expose.

@mcp.tool()
def create_word_document(filepath: str, content: str) -> Dict[str, Any]:
    """
    Create a new Microsoft Word document with the provided content.
    
    Args:
        filepath: Path where to save the document
        content: Text content for the document
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Create a new document
        doc = Document()
        
        # Add content
        doc.add_paragraph(content)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        
        # Save the document
        doc.save(filepath)
        
        logger.info(f"Created Word document: {filepath}")
        return {
            "success": True,
            "message": "Successfully created Word document",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        return {
            "success": False,
            "message": f"Error creating Word document: {str(e)}",
            "filepath": None
        }

@mcp.tool()
def edit_word_document(filepath: str, operations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Edit an existing Microsoft Word document using the specified operations.
    
    Args:
        filepath: Path to the Word document
        operations: List of operations to perform, where each operation is a dictionary with:
            - type: Operation type (add_paragraph, add_heading, edit_paragraph, delete_paragraph)
            - Additional parameters depending on the operation type
            
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Load the document
        if not os.path.exists(filepath):
            return {
                "success": False,
                "message": f"File not found: {filepath}",
                "filepath": None
            }
        
        doc = Document(filepath)
        
        # Apply operations
        for op in operations:
            op_type = op.get("type")
            
            if op_type == "add_paragraph":
                doc.add_paragraph(op.get("text", ""))
            
            elif op_type == "add_heading":
                doc.add_heading(op.get("text", ""), level=op.get("level", 1))
            
            elif op_type == "edit_paragraph":
                idx = op.get("index", 0)
                new_text = op.get("text", "")
                
                if 0 <= idx < len(doc.paragraphs):
                    doc.paragraphs[idx].text = new_text
                else:
                    logger.warning(f"Paragraph index out of range: {idx}")
            
            elif op_type == "delete_paragraph":
                idx = op.get("index", 0)
                
                if 0 <= idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    p_elem = p._element
                    p_elem.getparent().remove(p_elem)
                else:
                    logger.warning(f"Paragraph index out of range: {idx}")
            
            else:
                logger.warning(f"Unknown operation type: {op_type}")
        
        # Save the document
        doc.save(filepath)
        
        logger.info(f"Edited Word document: {filepath}")
        return {
            "success": True,
            "message": "Successfully edited Word document",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error editing Word document: {str(e)}")
        return {
            "success": False,
            "message": f"Error editing Word document: {str(e)}",
            "filepath": None
        }

@mcp.tool()
def convert_txt_to_word(source_path: str, target_path: str) -> Dict[str, Any]:
    """
    Convert a text file to a Microsoft Word document.
    
    Args:
        source_path: Path to the text file
        target_path: Path where to save the Word document
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if source file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Source file not found: {source_path}",
                "filepath": None
            }
        
        # Read the text file
        with open(source_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        
        # Create a new document
        doc = Document()
        
        # Add content as paragraphs (split by newlines)
        for paragraph in text_content.split('\n'):
            if paragraph.strip():  # Skip empty paragraphs
                doc.add_paragraph(paragraph)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        
        # Save the document
        doc.save(target_path)
        
        logger.info(f"Converted text to Word: {source_path} -> {target_path}")
        return {
            "success": True,
            "message": "Successfully converted text to Word document",
            "filepath": target_path
        }
    except Exception as e:
        logger.error(f"Error converting text to Word: {str(e)}")
        return {
            "success": False,
            "message": f"Error converting text to Word: {str(e)}",
            "filepath": None
        }

# ---- Excel Operations ----

@mcp.tool()
def create_excel_file(filepath: str, content: str) -> Dict[str, Any]:
    """
    Create a new Excel file with the provided content.
    
    Args:
        filepath: Path where to save the Excel file
        content: Data content, either JSON string or CSV-like string
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Parse the content as JSON data
        try:
            data = json.loads(content)
        except json.JSONDecodeError:
            # If not valid JSON, treat as CSV
            data = [line.split(',') for line in content.strip().split('\n')]
        
        # Convert to DataFrame
        df = pd.DataFrame(data)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        
        # Save to Excel
        df.to_excel(filepath, index=False)
        
        logger.info(f"Created Excel file: {filepath}")
        return {
            "success": True,
            "message": "Successfully created Excel file",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error creating Excel file: {str(e)}")
        return {
            "success": False,
            "message": f"Error creating Excel file: {str(e)}",
            "filepath": None
        }

@mcp.tool()
def edit_excel_file(filepath: str, operations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Edit an existing Excel file using the specified operations.
    
    Args:
        filepath: Path to the Excel file
        operations: List of operations to perform, where each operation is a dictionary with:
            - type: Operation type (update_cell, update_range, delete_row, delete_column, add_sheet, delete_sheet)
            - Additional parameters depending on the operation type
            
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if file exists
        if not os.path.exists(filepath):
            return {
                "success": False,
                "message": f"File not found: {filepath}",
                "filepath": None
            }
        
        # Load the Excel file
        wb = openpyxl.load_workbook(filepath)
        
        # Apply operations
        for op in operations:
            op_type = op.get("type")
            sheet_name = op.get("sheet", wb.sheetnames[0])
            
            # Get the sheet, create if it doesn't exist
            if sheet_name not in wb.sheetnames:
                wb.create_sheet(sheet_name)
            
            sheet = wb[sheet_name]
            
            if op_type == "update_cell":
                row = op.get("row", 1)
                col = op.get("col", 1)
                value = op.get("value", "")
                
                sheet.cell(row=row, column=col, value=value)
            
            elif op_type == "update_range":
                start_row = op.get("start_row", 1)
                start_col = op.get("start_col", 1)
                values = op.get("values", [])
                
                for i, row_values in enumerate(values):
                    for j, value in enumerate(row_values):
                        sheet.cell(row=start_row + i, column=start_col + j, value=value)
            
            elif op_type == "delete_row":
                row = op.get("row", 1)
                sheet.delete_rows(row)
            
            elif op_type == "delete_column":
                col = op.get("col", 1)
                sheet.delete_cols(col)
            
            elif op_type == "add_sheet":
                new_sheet_name = op.get("name", "NewSheet")
                if new_sheet_name not in wb.sheetnames:
                    wb.create_sheet(new_sheet_name)
            
            elif op_type == "delete_sheet":
                if sheet_name in wb.sheetnames and len(wb.sheetnames) > 1:
                    del wb[sheet_name]
            
            else:
                logger.warning(f"Unknown operation type: {op_type}")
        
        # Save the workbook
        wb.save(filepath)
        
        logger.info(f"Edited Excel file: {filepath}")
        return {
            "success": True,
            "message": "Successfully edited Excel file",
            "filepath": filepath
        }
    except Exception as e:
        logger.error(f"Error editing Excel file: {str(e)}")
        return {
            "success": False,
            "message": f"Error editing Excel file: {str(e)}",
            "filepath": None
        }

@mcp.tool()
def convert_csv_to_excel(source_path: str, target_path: str) -> Dict[str, Any]:
    """
    Convert a CSV file to an Excel file.
    
    Args:
        source_path: Path to the CSV file
        target_path: Path where to save the Excel file
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if source file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Source file not found: {source_path}",
                "filepath": None
            }
        
        # Read the CSV file
        df = pd.read_csv(source_path)
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        
        # Save to Excel
        df.to_excel(target_path, index=False)
        
        logger.info(f"Converted CSV to Excel: {source_path} -> {target_path}")
        return {
            "success": True,
            "message": "Successfully converted CSV to Excel",
            "filepath": target_path
        }
    except Exception as e:
        logger.error(f"Error converting CSV to Excel: {str(e)}")
        return {
            "success": False,
            "message": f"Error converting CSV to Excel: {str(e)}",
            "filepath": None
        }

# ---- PDF Operations ----

# @mcp.tool()
# def create_pdf_file(filepath: str, content: str) -> Dict[str, Any]:
#     """
#     Create a new PDF file with the provided text content.
    
#     Args:
#         filepath: Path where to save the PDF file
#         content: Text content for the PDF
        
#     Returns:
#         Operation result with success status, message, and filepath
#     """
#     try:
#         # Ensure the directory exists
#         os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)
        
#         # Create a new PDF with ReportLab
#         c = canvas.Canvas(filepath, pagesize=letter)
#         width, height = letter
        
#         # Process text content
#         lines = content.split('\n')
        
#         y_position = height - 40  # Start position from top
#         for line in lines:
#             if y_position < 40:  # If we're at the bottom of the page
#                 c.showPage()  # Create a new page
#                 y_position = height - 40  # Reset position
            
#             c.drawString(40, y_position, line)
#             y_position -= 15  # Move down for next line
        
#         c.save()
        
#         logger.info(f"Created PDF file: {filepath}")
#         return {
#             "success": True,
#             "message": "Successfully created PDF file",
#             "filepath": filepath
#         }
#     except Exception as e:
#         logger.error(f"Error creating PDF file: {str(e)}")
#         return {
#             "success": False,
#             "message": f"Error creating PDF file: {str(e)}",
#             "filepath": None
#         }

@mcp.tool()
def convert_word_to_pdf(source_path: str, target_path: str) -> Dict[str, Any]:
    """
    Convert a Microsoft Word document to a PDF file.
    
    Args:
        source_path: Path to the Word document
        target_path: Path where to save the PDF file
        
    Returns:
        Operation result with success status, message, and filepath
    """
    try:
        # Check if source file exists
        if not os.path.exists(source_path):
            return {
                "success": False,
                "message": f"Source file not found: {source_path}",
                "filepath": None
            }
        
        # Ensure the directory exists
        os.makedirs(os.path.dirname(os.path.abspath(target_path)), exist_ok=True)
        
        # Convert Word to PDF using docx2pdf
        docx2pdf.convert(source_path, target_path)
        
        logger.info(f"Converted Word to PDF: {source_path} -> {target_path}")
        return {
            "success": True,
            "message": "Successfully converted Word to PDF",
            "filepath": target_path
        }
    except Exception as e:
        logger.error(f"Error converting Word to PDF: {str(e)}")
        return {
            "success": False,
            "message": f"Error converting Word to PDF: {str(e)}",
            "filepath": None
        }

@mcp.tool()
def generate_pdf_from_markdown(content):
    """
    Generates a PDF file from a markdown string using the HTML/CSS method.

    This function first converts the markdown content to an HTML string,
    then uses WeasyPrint to render the HTML and CSS into a PDF document.

    Args:
        content (str): The markdown string content.
    """
    
    # 1. Convert Markdown to HTML
    html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])

    # 2. Add inline CSS for styling the HTML
    # WeasyPrint works best with CSS for styling.
    css_string = """
    @page {
        size: A4;
        margin: 0.5in;
    }
    body {
        font-family: 'Helvetica', sans-serif;
        font-size: 12pt;
        line-height: 1.5;
        color: #333333;
    }
    h1 {
        font-size: 24pt;
        color: #2a4365;
        text-align: center;
        margin-top: 1em;
        margin-bottom: 0.5em;
        padding-bottom: 0.2em;
        border-bottom: 2px solid #e2e8f0;
    }
    h2 {
        font-size: 18pt;
        color: #4a5568;
        margin-top: 1.5em;
        margin-bottom: 0.5em;
    }
    h3 {
        font-size: 16pt;
        color: #718096;
        margin-top: 1.5em;
        margin-bottom: 0.5em;
    }
    h4 {
        font-size: 14pt;
        color: #a0aec0;
        margin-top: 1.5em;
        margin-bottom: 0.5em;
    }
    p {
        margin-top: 0.5em;
        margin-bottom: 1em;
        text-align: justify;
    }
    table {
        border-collapse: collapse;
        width: 100%;
        margin-top: 1em;
        margin-bottom: 1em;
    }
    th, td {
        border: 1px solid #e2e8f0;
        padding: 8px;
        text-align: left;
    }
    th {
        background-color: #4a5568;
        color: white;
        font-weight: bold;
    }
    tr:nth-child(even) {
        background-color: #f7fafc;
    }
    code {
        font-family: 'Courier New', Courier, monospace;
        background-color: #e2e8f0;
        padding: 2px 4px;
        border-radius: 4px;
    }
    pre {
        background-color: #e2e8f0;
        padding: 10px;
        border-radius: 8px;
        overflow-x: auto;
    }
    """
    
    # Create the complete HTML document
    full_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Generated PDF</title>
        <style>{css_string}</style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    # 3. Use WeasyPrint to render the HTML to PDF
    # Configure fonts to avoid rendering issues
    font_config = FontConfiguration()
    
    # Create an HTML object and render it to a PDF
    html = HTML(string=full_html, base_url='.')
    pdf_bytes = html.write_pdf(stylesheets=[CSS(string=css_string, font_config=font_config)])

    # 4. Save the PDF to a file
    with open("output.pdf", 'wb') as f:
        f.write(pdf_bytes)

    print(f"PDF file output.pdf generated successfully.")


if __name__ == "__main__":
    logger.info(f"🚀 MCP server started on port {os.getenv('PORT', 8084)}")
    # Could also use 'sse' transport, host="0.0.0.0" required for Cloud Run.
    asyncio.run(
        mcp.run_async(
            transport="streamable-http",
            host="0.0.0.0",
            port=os.getenv("PORT", 8084),
        )
    )
# --- End MCP Server ---