

import json
import os

from pathlib import Path
from typing import Dict, Any, List, Optional
from dotenv import load_dotenv
import io
import re
from bs4 import BeautifulSoup
import google.genai.types as types

import markdown
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration
from lxml import etree
from google.adk.tools import ToolContext

# --- Load Environment Variables (If ADK tools need them, e.g., API keys) ---
load_dotenv() # Create a .env file in the same directory if needed

# Document processing libraries
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    raise ImportError("Please install python-docx with: uv pip install python-docx")

try:
    import pandas as pd
    import openpyxl
except ImportError:
    raise ImportError("Please install pandas and openpyxl with: uv pip install pandas openpyxl")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    raise ImportError("Please install reportlab with: uv pip install reportlab")

try:
    import docx2pdf
except ImportError:
    raise ImportError("Please install docx2pdf with: uv pip install docx2pdf")

async def generate_docx_from_markdown(content: str, title: str,  doc_type: str, tool_context: ToolContext) -> dict:
    """
    Generates a Word document (.docx) from a markdown string, creates an ADK artifact,
    saves it, and returns a JSON object indicating the status of the operation.

    Args:
        content (str): The markdown content to be converted.
        title (str): The title for the document's cover page. This can be the request user look for or a suitable title you can suggest for the document.
        doc_type (str): The type of document, used for naming the file. It can be either "proposal", "report" or "document" depends to the content input. Defaults to "document".

    Returns:
        dict: A JSON object with a 'status' key: {"status": "success"} or {"status": "failed"}.
    """
    try:
        document = Document()
        document.add_heading(title, level=0)
        document.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_page_break()

        html = markdown.markdown(content, extensions=['tables', 'fenced_code', 'extra'])
        html = markdown.markdown(content, extensions=['tables', 'fenced_code', 'extra'])

        def create_hyperlink(paragraph, url, text):
            part = paragraph.part
            r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rStyle = OxmlElement('w:rStyle')
            rStyle.set(qn('w:val'), 'Hyperlink')
            rPr.append(rStyle)
            new_run.append(rPr)
            new_run.append(OxmlElement('w:t'))
            new_run.text = text
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

        def add_formatted_run(paragraph, text):
            inline_tags_pattern = re.compile(r'<strong>(.*?)</strong>|<em>(.*?)</em>|<code>(.*?)</code>|<a\s+href="([^"]*)">(.*?)</a>', re.DOTALL)
            
            current_text = text
            while True:
                match = inline_tags_pattern.search(current_text)
                if not match:
                    if current_text:
                        paragraph.add_run(current_text)
                    break
                
                paragraph.add_run(current_text[:match.start()])
                
                if match.group(1):
                    paragraph.add_run(match.group(1)).bold = True
                elif match.group(2):
                    paragraph.add_run(match.group(2)).italic = True
                elif match.group(3):
                    run = paragraph.add_run(match.group(3))
                    run.font.name = 'Courier New'
                elif match.group(4) and match.group(5):
                    create_hyperlink(paragraph, match.group(4), match.group(5))
                
                current_text = current_text[match.end():]

        def parse_table_and_add(html_block, document):
            rows = re.findall(r'<tr.*?>(.*?)</tr>', html_block, re.DOTALL)
            if not rows:
                return

            # Find all th/td tags to get the maximum number of columns
            all_cells = re.findall(r'<t[dh].*?>', html_block, re.DOTALL)
            if not all_cells:
                return
            
            # The number of columns is the max number of cells in any row
            num_cols = max(len(re.findall(r'<t[dh].*?>', row)) for row in rows)
            table = document.add_table(rows=len(rows), cols=num_cols)
            table.autofit = True
            
            for i, row in enumerate(rows):
                cells = re.findall(r'<t[dh].*?>(.*?)</t[dh]>', row, re.DOTALL)
                for j, cell_text in enumerate(cells):
                    paragraph = table.cell(i, j).paragraphs[0]
                    add_formatted_run(paragraph, cell_text.strip())
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        def parse_list_and_add(html_list_block, document, indent_level=0, list_style='List Bullet'):
            # Regex to find all list items and nested lists at once
            list_items = re.findall(r'<li>(.*?)</li>', html_list_block, re.DOTALL)
            
            for item_content in list_items:
                # Add the list item paragraph
                p = document.add_paragraph(style=list_style)
                p.paragraph_format.left_indent = Inches(0.25 * indent_level)
                
                # Check for nested lists within the item
                nested_list_match = re.search(r'<ul.*?</ul>|<ol.*?<ol>', item_content, re.DOTALL)
                
                if nested_list_match:
                    # Content before the nested list
                    text_before = item_content[:nested_list_match.start()]
                    add_formatted_run(p, text_before.strip())
                    
                    # Recursively parse the nested list
                    parse_list_and_add(nested_list_match.group(0), document, indent_level + 1, list_style)
                    
                    # Content after the nested list
                    text_after = item_content[nested_list_match.end():]
                    if text_after.strip():
                        p_after = document.add_paragraph(style=list_style)
                        p_after.paragraph_format.left_indent = Inches(0.25 * indent_level)
                        add_formatted_run(p_after, text_after.strip())
                else:
                    add_formatted_run(p, item_content.strip())

        def parse_html_content(html_content):
            # This regex matches the entire block for each type of tag
            blocks_pattern = re.compile(
                r'(<h[1-6]>.*?</h[1-6]>|<p>.*?</p>|<ul.*?</ul.*?>|<ol.*?</ol.*?>|<pre>.*?</pre>|<br\s*/>|<table.*?</table.*?>)',
                re.DOTALL
            )
            
            last_end = 0
            
            for match in blocks_pattern.finditer(html_content):
                text_before = html_content[last_end:match.start()].strip()
                if text_before:
                    p = document.add_paragraph()
                    add_formatted_run(p, text_before)

                block = match.group(0).strip()
                
                if block.startswith('<h'):
                    heading_match = re.match(r'<h(\d)>(.*?)</h\1>', block, re.DOTALL)
                    if heading_match:
                        level = int(heading_match.group(1))
                        text = heading_match.group(2).strip()
                        document.add_heading(text, level=level)
                elif block.startswith('<p>'):
                    p = document.add_paragraph()
                    text = re.sub(r'</?p>', '', block).strip()
                    add_formatted_run(p, text)
                elif block.startswith('<ul') or block.startswith('<ol'):
                    list_style = 'List Bullet' if block.startswith('<ul') else 'List Number'
                    parse_list_and_add(block, document, 0, list_style)
                elif block.startswith('<table'):
                    parse_table_and_add(block, document)
                elif block.startswith('<pre>'):
                    code_text_match = re.search(r'<code>(.*?)</code>', block, re.DOTALL)
                    if code_text_match:
                        text = code_text_match.group(1)
                        document.add_paragraph(text, style='No Spacing')
                
                last_end = match.end()

            remaining_text = html_content[last_end:].strip()
            if remaining_text:
                p = document.add_paragraph()
                add_formatted_run(p, remaining_text)
        
        parse_html_content(html)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        
        docx_bytes = buffer.getvalue()

        session_id = tool_context._invocation_context.session.id
        unique_id = os.urandom(4).hex()
        output_filename = f"{session_id}_{unique_id}_{doc_type.lower()}.docx"
        
        # Create the ADK artifact with the correct MIME type
        artifact = types.Part(
            inline_data=types.Blob(
                data=docx_bytes,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        )
        
        await tool_context.save_artifact(
            filename=output_filename,
            artifact=artifact
        )
        
        return {"status": "success", "filename": output_filename}
        
    except Exception as e:
        print(f"An error occurred: {e}")
        return {"status": "failed", "error": str(e)}

async def generate_pdf_from_markdown(content:str, doc_type:str, tool_context: ToolContext) -> dict:
    """
    Generates a PDF file from a markdown string using the HTML/CSS method.

    This function first converts the markdown content to an HTML string,
    then uses WeasyPrint to render the HTML and CSS into a PDF document.

    Args:
        content (str): The markdown string content.
        doc_type (str): The type of document, used for naming the file. It can be either "proposal", "report" or "document" depends to the content input. Defaults to "document".
    """
    try: 
        # 1. Convert Markdown to HTML
        html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])

        # 2. Add inline CSS for styling the HTML
        # WeasyPrint works best with CSS for styling.
        css_string = """
        @page {
            size: A4;
            margin: 0.5in;
        }
        body {
            font-family: 'Helvetica', sans-serif;
            font-size: 12pt;
            line-height: 1.5;
            color: #333333;
        }
        h1 {
            font-size: 24pt;
            color: #2a4365;
            text-align: center;
            margin-top: 1em;
            margin-bottom: 0.5em;
            padding-bottom: 0.2em;
            border-bottom: 2px solid #e2e8f0;
        }
        h2 {
            font-size: 18pt;
            color: #4a5568;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        h3 {
            font-size: 16pt;
            color: #718096;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        h4 {
            font-size: 14pt;
            color: #a0aec0;
            margin-top: 1.5em;
            margin-bottom: 0.5em;
        }
        p {
            margin-top: 0.5em;
            margin-bottom: 1em;
            text-align: justify;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin-top: 1em;
            margin-bottom: 1em;
        }
        th, td {
            border: 1px solid #e2e8f0;
            padding: 8px;
            text-align: left;
        }
        th {
            background-color: #4a5568;
            color: white;
            font-weight: bold;
        }
        tr:nth-child(even) {
            background-color: #f7fafc;
        }
        code {
            font-family: 'Courier New', Courier, monospace;
            background-color: #e2e8f0;
            padding: 2px 4px;
            border-radius: 4px;
        }
        pre {
            background-color: #e2e8f0;
            padding: 10px;
            border-radius: 8px;
            overflow-x: auto;
        }
        """
        
        # Create the complete HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Generated PDF</title>
            <style>{css_string}</style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        # 3. Use WeasyPrint to render the HTML to PDF
        # Configure fonts to avoid rendering issues
        font_config = FontConfiguration()
        
        # Create an HTML object and render it to a PDF
        html = HTML(string=full_html, base_url='.')
        pdf_bytes = html.write_pdf(stylesheets=[CSS(string=css_string, font_config=font_config)])

        session_id = tool_context._invocation_context.session.id
        unique_id = os.urandom(4).hex()
        output_filename = f"{session_id}_{unique_id}_{doc_type.lower()}.pdf"
            
            # Create the ADK artifact with the correct MIME type
        artifact = types.Part(
            inline_data=types.Blob(
                data=pdf_bytes,
                mime_type="application/pdf"
            )
        )
        
        await tool_context.save_artifact(
            filename=output_filename,
            artifact=artifact
        )
        
        return {"status": "success", "filename": output_filename}

    except Exception as e:
        print(f"An error occurred: {e}")
        return {"status": "failed", "error": str(e)}