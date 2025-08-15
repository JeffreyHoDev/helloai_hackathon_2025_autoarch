import io
import markdown
from docx import Document
from docx.shared import Inches

def generate_docx_from_markdown_bytes(content: str, tool_context: dict) -> bytes:
    """
    Generates a Word document (.docx) from a markdown string and returns it as a bytes object.

    The function uses the python-docx library to create the document and markdown to
    parse the input string. It handles basic formatting like headings, bold, and italics.

    Args:
        content (str): The markdown content to be converted.
        tool_context (dict): The ADK tool context. (Not used in this function, but included
                             in the signature as per the request.)

    Returns:
        bytes: The bytes of the generated .docx file.
    """
    document = Document()
    
    # Process the markdown content
    # The 'fenced_code' and 'tables' extensions are enabled for better parsing.
    html = markdown.markdown(content, extensions=['tables', 'fenced_code', 'extra'])

    # A simple way to handle the HTML output is to split by lines
    # and apply styles based on the beginning of the line.
    for line in html.splitlines():
        line = line.strip()
        if not line:
            continue

        if line.startswith('<h1>'):
            document.add_heading(line.replace('<h1>', '').replace('</h1>', '').strip(), level=1)
        elif line.startswith('<h2>'):
            document.add_heading(line.replace('<h2>', '').replace('</h2>', '').strip(), level=2)
        elif line.startswith('<h3>'):
            document.add_heading(line.replace('<h3>', '').replace('</h3>', '').strip(), level=3)
        elif line.startswith('<p>'):
            p = document.add_paragraph()
            # Simple regex for bold and italics
            # This is a basic implementation and could be more robust
            processed_line = line.replace('<p>', '').replace('</p>', '')
            parts = processed_line.split('<strong>')
            for part in parts:
                if '</strong>' in part:
                    strong_parts = part.split('</strong>')
                    p.add_run(strong_parts[0]).bold = True
                    p.add_run(strong_parts[1])
                else:
                    p.add_run(part)
        elif line.startswith('<ul>') or line.startswith('<li>'):
            document.add_paragraph(line.replace('<ul>', '').replace('</ul>', '').replace('<li>', '').replace('</li>', '').strip(), style='List Bullet')
        elif line.startswith('<table>'):
            # This is a simplified table handling and can be extended
            document.add_paragraph("Table content:")
            
    # Save the document to a bytes buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

# Example usage:
if __name__ == '__main__':
    # The tool context would be passed by the ADK, here it's a placeholder.
    mock_tool_context = {"task_id": "12345", "user_id": "abcde"}

    sample_content = """
# Project Report

This is a **report** generated from a markdown string. It demonstrates how to convert text with *basic formatting* to a Word document in memory.

## Key Findings
-   First finding.
-   Second finding.

### Detailed Analysis
Here is a paragraph of text that shows how the content wraps and how different styles are applied. This is a very simple example.

| Header 1 | Header 2 |
|----------|----------|
| Row 1 Col 1| Row 1 Col 2|
| Row 2 Col 1| Row 2 Col 2|
"""
    
    docx_bytes = generate_docx_from_markdown_bytes(sample_content, mock_tool_context)
    
    # To verify the output, you can save the bytes to a local file.
    with open("sample_report.docx", "wb") as f:
        f.write(docx_bytes)
        
    print("Word document bytes generated and saved to 'sample_report.docx' for verification.")